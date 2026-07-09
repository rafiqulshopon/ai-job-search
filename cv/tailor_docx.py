#!/usr/bin/env python3
"""tailor_docx.py - tailor a Google Docs (.docx) CV per job application.

Section-level operations that preserve Google Docs styling by editing runs
in-place and cloning existing list paragraphs (so bullets keep their list
style + formatting). The Edit tool cannot touch a binary .docx, so /apply
routes all CV edits through this helper.

Usage:
  python3 cv/tailor_docx.py read <path>
  python3 cv/tailor_docx.py copy <dst> [--from cv/main_example.docx]
  python3 cv/tailor_docx.py set-profile <path> --text "..."
  python3 cv/tailor_docx.py set-bullets <path> --role "<job title>" --bullets "b1|b2|b3"
  python3 cv/tailor_docx.py set-skills <path> --items "Frontend: ...|Backend: ...|Tools: ..."

Bullets and skills items are pipe (|) separated.
"""
import argparse
import copy
import os
import shutil
import sys

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph
from docx.text.run import Run

ICONS_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "_assets")

SECTION_HEADERS = ("SUMMARY", "PROFESSIONAL EXPERIENCE", "EXPERIENCE", "EDUCATION",
                   "SKILLS", "SKILLS & OTHER", "LANGUAGES", "CERTIFICATIONS")


# ---------- low-level helpers ----------

def _is_list(para):
    """True if the paragraph is a list item (has w:numPr)."""
    pPr = para._p.find(qn("w:pPr"))
    return pPr is not None and pPr.find(qn("w:numPr")) is not None


def _is_section_header(text):
    t = text.strip().rstrip(":").upper()
    return any(t == h or t.startswith(h) for h in SECTION_HEADERS) and len(t) < 40


def _set_para_text(para, text):
    """Replace paragraph text while preserving the first run's formatting.

    Also drops any w:hyperlink elements first: python-docx hides the runs nested
    inside a hyperlink from para.runs, so without this a rewritten bullet can be
    left with stale link display text appended (e.g. '...Pusher.property management')."""
    p = para._p
    for hl in p.findall(qn("w:hyperlink")):
        p.remove(hl)
    runs = para.runs
    if runs:
        runs[0].text = text
        for r in runs[1:]:
            r._element.getparent().remove(r._element)
    else:
        para.add_run(text)


def _set_skill_line(para, text):
    """Replace a SKILLS line keeping the category label bold and the content normal.

    A skill line is 'Category: content'. The master structures each one as a bold
    category run followed by one or more normal-weight runs. _set_para_text folds the
    whole line into the first (bold) run, making everything bold; this instead splits
    on the first colon so only the label stays bold. Font family/size are preserved
    because existing runs are edited in place (the first run cloned for the content
    when the line was a single all-bold run)."""
    p = para._p
    for hl in p.findall(qn("w:hyperlink")):
        p.remove(hl)
    runs = para.runs
    if not runs:
        para.add_run(text)
        return
    label, sep, after = text.partition(":")
    if not sep:
        # No colon: treat the whole line as the (bold) label.
        runs[0].text = text
        runs[0].bold = True
        for r in runs[1:]:
            r._element.getparent().remove(r._element)
        return
    label = label.rstrip()
    rest = ":" + after  # keep the original spacing after the colon
    # Bold category label in the first run.
    first = runs[0]
    first.text = label
    first.bold = True
    # Normal-weight content in the second run (clone the first if there is only one).
    if len(runs) >= 2:
        second = runs[1]
        second.text = rest
        second.bold = False
        for r in runs[2:]:
            r._element.getparent().remove(r._element)
    else:
        new_r = copy.deepcopy(first._element)
        first._element.addnext(new_r)
        second = Run(new_r, para)
        second.text = rest
        second.bold = False


def _run_hint(para):
    """Formatting hint for the first run (helps Claude see styling)."""
    if not para.runs:
        return ""
    r = para.runs[0]
    bits = []
    if r.bold:
        bits.append("bold")
    if r.italic:
        bits.append("italic")
    if r.font and r.font.name:
        bits.append(r.font.name)
    if r.font and r.font.size:
        bits.append(f"{r.font.size.pt}pt")
    return ",".join(bits)


def _find_paragraph(paras, needle):
    """Find first paragraph whose text matches `needle` (exact, then contains). Case-insensitive."""
    needle_l = needle.strip().lower()
    for i, p in enumerate(paras):
        if p.text.strip().lower() == needle_l:
            return i
    for i, p in enumerate(paras):
        if needle_l and needle_l in p.text.strip().lower():
            return i
    return -1


def _find_section(paras, name):
    """Find a section header paragraph (e.g. SUMMARY). Returns its index or -1."""
    name_l = name.strip().lower()
    for i, p in enumerate(paras):
        if p.text.strip().rstrip(":").lower() == name_l:
            return i
    for i, p in enumerate(paras):
        t = p.text.strip().rstrip(":").lower()
        if name_l in t and len(t) < 40:
            return i
    return -1


def _delete_paragraph(para):
    para._element.getparent().remove(para._element)


def _clone_after(template_para, after_para, text, setter=_set_para_text):
    """Clone template_para's <w:p> (keeps list style + formatting), insert after after_para, set text."""
    new_p = copy.deepcopy(template_para._p)
    after_para._p.addnext(new_p)
    new_para = Paragraph(new_p, template_para._parent)
    setter(new_para, text)
    return new_para


def _replace_block(paras, template_para, old_block, new_texts, start_anchor, setter=_set_para_text):
    """Replace old_block (list of paragraphs) with new_texts, cloning template_para for extras.

    start_anchor is the paragraph before old_block (insertion point for clones)."""
    # 1. Set text on existing block paragraphs (up to min length)
    n_old, n_new = len(old_block), len(new_texts)
    for i in range(min(n_old, n_new)):
        setter(old_block[i], new_texts[i])
    # 2. If fewer new texts -> delete extra old paragraphs
    if n_new < n_old:
        for p in old_block[n_new:]:
            _delete_paragraph(p)
    # 3. If more new texts -> clone template after the last kept paragraph
    elif n_new > n_old:
        anchor = old_block[-1] if old_block else start_anchor
        for extra in new_texts[n_old:]:
            anchor = _clone_after(template_para, anchor, extra, setter=setter)


# ---------- block collectors ----------

def _collect_bullets_after(paras, idx):
    """Consecutive list paragraphs after idx; tolerates one leading non-list (job title) + empties."""
    bullets, started, skipped = [], False, 0
    j = idx + 1
    while j < len(paras):
        p = paras[j]
        if _is_list(p):
            bullets.append(p); started = True
        elif not p.text.strip():
            pass  # skip blank
        else:
            if not started and skipped < 1:
                skipped += 1  # tolerate a single title line between company and bullets
            else:
                break
        j += 1
    return bullets


def _collect_section_body(paras, header_idx):
    """Non-empty paragraphs after a section header until the next header / end."""
    body = []
    for j in range(header_idx + 1, len(paras)):
        t = paras[j].text.strip()
        if not t:
            continue
        if _is_section_header(t):
            break
        body.append(paras[j])
    return body


# ---------- single-page cleanup ----------

def _neutralize_section_breaks(doc):
    """Set inline (mid-document) section breaks to 'continuous' so they don't force a new page."""
    n = 0
    for pPr in doc.element.body.iter(qn("w:pPr")):
        sect = pPr.find(qn("w:sectPr"))
        if sect is not None:
            typ = sect.find(qn("w:type"))
            if typ is None:
                typ = OxmlElement("w:type")
                typ.set(qn("w:val"), "continuous")
                sect.insert(0, typ)
            else:
                typ.set(qn("w:val"), "continuous")
            n += 1
    return n


def _strip_trailing_empty(doc):
    """Remove trailing empty paragraphs (skip the one carrying the final sectPr)."""
    removed = 0
    paras = doc.paragraphs
    while paras and not paras[-1].text.strip():
        el = paras[-1]._element
        pPr = el.find(qn("w:pPr"))
        if pPr is not None and pPr.find(qn("w:sectPr")) is not None:
            break
        el.getparent().remove(el)
        paras = doc.paragraphs
        removed += 1
    return removed


def _reduce_vertical_margins(doc, by_cm=0.4, floor_cm=0.8):
    """Shrink top/bottom margins to give a buffer against a blank trailing page (Google Docs vs
    LibreOffice rendering often leaves content exactly filling page 1)."""
    touched = 0
    for s in doc.sections:
        if s.top_margin:
            s.top_margin = Cm(max(s.top_margin.cm - by_cm, floor_cm))
        if s.bottom_margin:
            s.bottom_margin = Cm(max(s.bottom_margin.cm - by_cm, floor_cm))
        touched += 1
    return touched


def clean_for_single_page(doc):
    """Apply all three blank-trailing-page fixes. Returns a summary string."""
    n = _neutralize_section_breaks(doc)
    r = _strip_trailing_empty(doc)
    m = _reduce_vertical_margins(doc)
    return f"sections->continuous:{n} trailing-empty-removed:{r} margins-reduced-in:{m}"


# ---------- commands ----------

def cmd_read(args):
    doc = Document(args.path)
    paras = doc.paragraphs
    print(f"# {args.path}  ({len(paras)} paragraphs, {len(doc.tables)} tables)")
    for i, p in enumerate(paras):
        t = p.text.strip()
        marker = "bullet " if _is_list(p) else ""
        hint = _run_hint(p)
        hint_s = f" [{hint}]" if hint else ""
        preview = t[:95] + ("..." if len(t) > 95 else "")
        print(f"[{i:02d}] {marker}{preview}{hint_s}")


def cmd_copy(args):
    src = args.src
    if not os.path.exists(src):
        sys.exit(f"source not found: {src}")
    shutil.copy2(src, args.dst)  # copy2 keeps file metadata
    # Apply single-page cleanup (Google Docs exports often carry a trailing blank page)
    doc = Document(args.dst)
    summary = clean_for_single_page(doc)
    doc.save(args.dst)
    print(f"copied {src} -> {args.dst}  [single-page cleanup: {summary}]")


def cmd_tighten_margins(args):
    doc = Document(args.path)
    m = _reduce_vertical_margins(doc, by_cm=args.by, floor_cm=0.8)
    doc.save(args.path)
    print(f"reduced top/bottom margins by {args.by}cm across {m} section(s)")


def cmd_set_profile(args):
    doc = Document(args.path)
    paras = doc.paragraphs
    idx = _find_section(paras, "SUMMARY")
    if idx < 0:
        sys.exit("SUMMARY section not found")
    # first non-empty paragraph after the header is the profile text
    for j in range(idx + 1, len(paras)):
        if paras[j].text.strip():
            _set_para_text(paras[j], args.text)
            doc.save(args.path)
            print(f"set profile text ({len(args.text)} chars)")
            return
    sys.exit("no profile paragraph found under SUMMARY")


def cmd_set_bullets(args):
    doc = Document(args.path)
    paras = doc.paragraphs
    idx = _find_paragraph(paras, args.role)
    if idx < 0:
        sys.exit(f"role not found: {args.role}")
    bullets = _collect_bullets_after(paras, idx)
    if not bullets:
        sys.exit(f"no bullets found after role '{args.role}'")
    new_texts = [b.strip() for b in args.bullets.split("|") if b.strip()]
    anchor = paras[idx]
    _replace_block(paras, bullets[0], bullets, new_texts, anchor)
    doc.save(args.path)
    print(f"set {len(new_texts)} bullets for '{args.role}' (was {len(bullets)})")


def cmd_set_skills(args):
    doc = Document(args.path)
    paras = doc.paragraphs
    idx = _find_section(paras, "SKILLS")
    if idx < 0:
        idx = _find_section(paras, "SKILLS & OTHER")
    if idx < 0:
        sys.exit("SKILLS section not found")
    body = _collect_section_body(paras, idx)
    if not body:
        sys.exit("no skills lines found under SKILLS")
    new_texts = [s.strip() for s in args.items.split("|") if s.strip()]
    _replace_block(paras, body[0], body, new_texts, paras[idx], setter=_set_skill_line)
    doc.save(args.path)
    print(f"set {len(new_texts)} skills lines (was {len(body)})")


# ---------- contact line (icons + clickable usernames) ----------

def _clear_paragraph_body(para):
    """Remove all runs and hyperlinks from a paragraph, preserving its w:pPr (alignment etc.)."""
    p = para._p
    for child in list(p):
        if child.tag in (qn("w:r"), qn("w:hyperlink")):
            p.remove(child)


def _style_run(run, font_name, size_pt):
    run.font.name = font_name
    if size_pt:
        run.font.size = Pt(size_pt)


def _add_hyperlink(para, url, text, font_name, size_pt, color="000000", underline=False):
    """Append a hyperlinked, styled run (black, no underline by default - clean for a resume)."""
    r_id = para.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rFonts.set(qn(attr), font_name)
    rPr.append(rFonts)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(int(size_pt * 2))); rPr.append(sz)
    szCs = OxmlElement("w:szCs"); szCs.set(qn("w:val"), str(int(size_pt * 2))); rPr.append(szCs)
    col = OxmlElement("w:color"); col.set(qn("w:val"), color); rPr.append(col)
    if underline:
        u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    r.append(rPr)
    t = OxmlElement("w:t"); t.text = text; t.set(qn("xml:space"), "preserve")
    r.append(t)
    hyperlink.append(r)
    para._p.append(hyperlink)


def _add_icon(para, path, height_pt, font_name):
    run = para.add_run()
    _style_run(run, font_name, height_pt)  # keep font for fallback; image ignores font/size
    run.add_picture(path, height=Pt(height_pt))
    return run


def cmd_set_contact(args):
    doc = Document(args.path)
    paras = doc.paragraphs
    idx = next((i for i, p in enumerate(paras) if "@" in p.text), -1)
    if idx < 0:
        sys.exit("contact line (paragraph containing '@') not found")
    para = paras[idx]
    orig = para.runs[0] if para.runs else None
    fname = (orig.font.name if orig and orig.font.name else "Garamond")
    size_pt = (orig.font.size.pt if orig and orig.font.size else 10)
    li_icon = args.linkedin_icon or os.path.join(ICONS_DIR, "linkedin.png")
    gh_icon = args.github_icon or os.path.join(ICONS_DIR, "github.png")
    for icon in (li_icon, gh_icon):
        if not os.path.exists(icon):
            sys.exit(f"icon not found: {icon} (place logos in cv/_assets/ or pass --linkedin-icon/--github-icon)")
    li_user = args.linkedin.rstrip("/").split("/")[-1]
    gh_user = args.github.rstrip("/").split("/")[-1]
    _clear_paragraph_body(para)
    lead = f"{args.location} • {args.phone} • {args.email} • "
    _style_run(para.add_run(lead), fname, size_pt)
    _add_icon(para, li_icon, size_pt, fname)
    _add_hyperlink(para, args.linkedin, f" {li_user}", fname, size_pt)
    _style_run(para.add_run(" • "), fname, size_pt)
    _add_icon(para, gh_icon, size_pt, fname)
    _add_hyperlink(para, args.github, f" {gh_user}", fname, size_pt)
    doc.save(args.path)
    print(f"rebuilt contact line with icons + clickable usernames (linkedin:{li_user}, github:{gh_user})")


# ---------- CLI ----------

def main():
    ap = argparse.ArgumentParser(description="Tailor a Google Docs (.docx) CV per application.")
    sub = ap.add_subparsers(dest="cmd", required=True)

    r = sub.add_parser("read", help="dump section structure")
    r.add_argument("path")
    r.set_defaults(func=cmd_read)

    c = sub.add_parser("copy", help="copy master to a tailored output")
    c.add_argument("dst")
    c.add_argument("--from", dest="src", default="cv/main_example.docx")
    c.set_defaults(func=cmd_copy)

    p = sub.add_parser("set-profile", help="replace the SUMMARY profile text")
    p.add_argument("path")
    p.add_argument("--text", required=True)
    p.set_defaults(func=cmd_set_profile)

    b = sub.add_parser("set-bullets", help="replace a role's experience bullets")
    b.add_argument("path")
    b.add_argument("--role", required=True, help="job title (e.g. 'Senior Frontend Engineer')")
    b.add_argument("--bullets", required=True, help="pipe-separated bullets")
    b.set_defaults(func=cmd_set_bullets)

    s = sub.add_parser("set-skills", help="replace the SKILLS lines")
    s.add_argument("path")
    s.add_argument("--items", required=True, help="pipe-separated skill lines")
    s.set_defaults(func=cmd_set_skills)

    tm = sub.add_parser("tighten-margins", help="reduce top/bottom margins to recover a blank trailing page")
    tm.add_argument("path")
    tm.add_argument("--by", type=float, default=0.4)
    tm.set_defaults(func=cmd_tighten_margins)

    sc = sub.add_parser("set-contact", help="rebuild the contact line with LinkedIn/GitHub icons + clickable usernames")
    sc.add_argument("path")
    sc.add_argument("--location", default="Dhaka, Bangladesh")
    sc.add_argument("--phone", default="+880 1991 895014")
    sc.add_argument("--email", default="rafiqul.dev@gmail.com")
    sc.add_argument("--linkedin", default="https://linkedin.com/in/rafiqulshopon")
    sc.add_argument("--github", default="https://github.com/rafiqulshopon")
    sc.add_argument("--linkedin-icon")
    sc.add_argument("--github-icon")
    sc.set_defaults(func=cmd_set_contact)

    args = ap.parse_args()
    args.func(args)


if __name__ == "__main__":
    main()
